"""Synchronize section 3.3.3 prose with the supplied Mushroomie ASM tone."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Inches
from docx.text.paragraph import Paragraph

try:
    from tools.asm_333_rewrites import PARAGRAPH_REWRITES, TABLE_CELL_REWRITES
except ModuleNotFoundError:  # Support direct execution from the tools directory.
    from asm_333_rewrites import PARAGRAPH_REWRITES, TABLE_CELL_REWRITES


SOURCE_PATH = Path(
    r"C:\Users\Admin\OneDrive\Tài liệu\mushroomie\artifacts"
    r"\Muc_3.3.3_Kenh_Website_Mushroomie_Khong_Landing_Page_SEO_Thuc_Te.docx"
)
OUTPUT_PATH = Path(
    r"C:\Users\Admin\OneDrive\Tài liệu\mushroomie\artifacts"
    r"\Muc_3.3.3_Kenh_Website_Mushroomie_Dong_Bo_Van_Phong_ASM.docx"
)

# Compatibility aliases retained for scripts written against the Task 2 brief.
SOURCE_DOCX = SOURCE_PATH
OUTPUT_DOCX = OUTPUT_PATH

NUMBER_PATTERN = re.compile(r"\b\d+(?:[.,]\d+)?(?:/\d+)?\b")


def _replace_text_preserving_format(paragraph: Paragraph, value: str) -> None:
    """Replace prose while retaining the paragraph and its leading run formatting."""
    if not paragraph.runs:
        paragraph.add_run(value)
        return

    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def _validate_rewrite_evidence() -> None:
    """Reject mappings that alter numeric evidence or introduce forbidden wording."""
    all_rewrites = PARAGRAPH_REWRITES | TABLE_CELL_REWRITES
    for old_text, new_text in all_rewrites.items():
        if NUMBER_PATTERN.findall(old_text) != NUMBER_PATTERN.findall(new_text):
            raise ValueError(f"Numeric evidence changed in rewrite: {old_text[:80]!r}")
        if "landing page" in new_text.lower():
            raise ValueError(f"Forbidden wording introduced in rewrite: {new_text[:80]!r}")


def apply_rewrites(document: DocumentType) -> tuple[int, int]:
    """Apply reviewed paragraph/cell mappings and return exact change counts."""
    _validate_rewrite_evidence()
    paragraph_count = 0
    cell_count = 0
    matched_paragraphs: set[str] = set()
    matched_cells: set[str] = set()

    for paragraph in document.paragraphs:
        original_text = paragraph.text
        if original_text in PARAGRAPH_REWRITES:
            _replace_text_preserving_format(
                paragraph, PARAGRAPH_REWRITES[original_text]
            )
            matched_paragraphs.add(original_text)
            paragraph_count += 1
        paragraph.paragraph_format.first_line_indent = Inches(0)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    if original_text in TABLE_CELL_REWRITES:
                        _replace_text_preserving_format(
                            paragraph, TABLE_CELL_REWRITES[original_text]
                        )
                        matched_cells.add(original_text)
                        cell_count += 1
                    paragraph.paragraph_format.first_line_indent = Inches(0)

    missing_paragraphs = set(PARAGRAPH_REWRITES) - matched_paragraphs
    missing_cells = set(TABLE_CELL_REWRITES) - matched_cells
    if missing_paragraphs or missing_cells:
        raise RuntimeError(
            "Reviewed mappings were not found in the source document: "
            f"paragraphs={len(missing_paragraphs)}, cells={len(missing_cells)}"
        )
    return paragraph_count, cell_count


def main() -> None:
    if not SOURCE_PATH.is_file():
        raise FileNotFoundError(f"Source document not found: {SOURCE_PATH}")

    document = Document(SOURCE_PATH)
    paragraph_count, cell_count = apply_rewrites(document)
    if paragraph_count < 30:
        raise RuntimeError(
            f"Số đoạn được biên tập chưa đủ phạm vi: {paragraph_count}"
        )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(
        f"paragraphs={paragraph_count}; table_cells={cell_count}; "
        f"output={OUTPUT_PATH}"
    )


if __name__ == "__main__":
    main()
