"""Audit the final Section 3.3.3 document for academic-tone regressions.

The audit is deliberately narrow: it checks editorial problems that can be
identified deterministically without changing the project's factual evidence.
It traverses both the document body and every table cell and reports stable
locations so an editor can correct a mapping rather than perform a broad,
untraceable replacement.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, Iterator

from docx import Document
from docx.document import Document as DocumentType


FORBIDDEN = {
    "first_person": ("chúng em", "chúng tôi", "em nhận thấy"),
    "promotional": ("hoàn hảo", "tốt nhất", "vượt trội", "chắc chắn"),
    "absolute_claim": ("hoàn toàn", "luôn luôn", "tuyệt đối"),
}

# These are formulaic transitions that add little analytical value when used
# repeatedly. A single purposeful use is acceptable; repeated use is not.
REPETITION_MARKERS = ("điều này cho thấy", "nhìn chung", "bên cạnh đó")

# The document retains established technical names (e.g. PageSpeed Insights,
# Lighthouse, CTA, KPI). This list contains English wording that has a direct,
# unambiguous Vietnamese equivalent and is not a product or metric name.
UNNECESSARY_ENGLISH = (
    "external link",
    "internal link",
    "mobile",
    "footer",
)

TERM_RULES: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("website_case", re.compile(r"\bwebsite\b")),
    ("duplicate_google_search_console", re.compile(r"Google\s+Google\s+Search\s+Console")),
    (
        "seo_onpage_case",
        re.compile(r"\bseo\s+onpage\b", re.IGNORECASE),
    ),
    (
        "pagespeed_insights_case",
        re.compile(r"\bpagespeed\s+insights\b", re.IGNORECASE),
    ),
    (
        "mushroomie_case",
        re.compile(r"(?<![\w.-])mushroomie(?![\w.-])", re.IGNORECASE),
    ),
)


@dataclass(frozen=True)
class TextRecord:
    location: str
    text: str


@dataclass(frozen=True)
class Finding:
    category: str
    location: str
    marker: str
    text: str


def iter_text_records(document: DocumentType) -> Iterator[TextRecord]:
    """Yield non-empty paragraphs in the body and all table cells."""
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            yield TextRecord(f"P{paragraph_index}", text)

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        yield TextRecord(
                            f"T{table_index}R{row_index}C{cell_index}P{paragraph_index}",
                            text,
                        )


def _find_phrase_matches(records: Iterable[TextRecord]) -> list[Finding]:
    findings: list[Finding] = []
    for record in records:
        lowered = record.text.casefold()
        for category, phrases in FORBIDDEN.items():
            for phrase in phrases:
                if phrase in lowered:
                    findings.append(Finding(category, record.location, phrase, record.text))

        for phrase in UNNECESSARY_ENGLISH:
            if phrase in lowered:
                findings.append(
                    Finding("terminology", record.location, phrase, record.text)
                )
    return findings


def _find_term_inconsistencies(records: Iterable[TextRecord]) -> list[Finding]:
    findings: list[Finding] = []
    for record in records:
        for marker, pattern in TERM_RULES:
            for match in pattern.finditer(record.text):
                value = match.group(0)
                if marker == "website_case" and value == "Website":
                    continue
                if marker == "seo_onpage_case" and value == "SEO Onpage":
                    continue
                if marker == "pagespeed_insights_case" and value == "PageSpeed Insights":
                    continue
                if marker == "mushroomie_case" and value == "Mushroomie":
                    continue
                findings.append(Finding("terminology", record.location, marker, record.text))
                break
    return findings


def _find_repetition(records: Iterable[TextRecord]) -> list[Finding]:
    findings: list[Finding] = []
    lowered_records = [(record, record.text.casefold()) for record in records]
    for marker in REPETITION_MARKERS:
        locations = [record for record, text in lowered_records if marker in text]
        if len(locations) > 2:
            for record in locations:
                findings.append(Finding("repetition", record.location, marker, record.text))
    return findings


def audit(records: Iterable[TextRecord]) -> list[Finding]:
    """Return stable, deduplicated tone findings for all supplied text."""
    record_list = list(records)
    findings = (
        _find_phrase_matches(record_list)
        + _find_term_inconsistencies(record_list)
        + _find_repetition(record_list)
    )
    return sorted(
        set(findings),
        key=lambda finding: (finding.category, finding.location, finding.marker),
    )


def _print_report(findings: list[Finding]) -> None:
    counts = Counter(finding.category for finding in findings)
    print(f"Tone audit: {len(findings)} finding(s)")
    for category in sorted(counts):
        print(f"- {category}: {counts[category]}")
        for finding in (item for item in findings if item.category == category):
            print(f"  {finding.location} [{finding.marker}]: {finding.text}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("document", type=Path, help="DOCX output to audit")
    parser.add_argument("--json", action="store_true", help="Print findings as JSON")
    args = parser.parse_args(argv)

    if not args.document.is_file():
        parser.error(f"Document not found: {args.document}")

    findings = audit(iter_text_records(Document(args.document)))
    if args.json:
        print(json.dumps([asdict(finding) for finding in findings], ensure_ascii=False))
    else:
        _print_report(findings)
    return 0 if not findings else 1


if __name__ == "__main__":
    sys.exit(main())
