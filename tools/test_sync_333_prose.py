"""Regression checks for the edited section 3.3.3 Word document."""

import hashlib
import re
import unittest
import zipfile

from docx import Document

from tools.asm_333_rewrites import PARAGRAPH_REWRITES, TABLE_CELL_REWRITES
from tools.sync_333_prose_to_asm import OUTPUT_PATH, SOURCE_PATH


def all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def document_text(document):
    return "\n".join(paragraph.text for paragraph in all_paragraphs(document))


def media_hashes(path):
    with zipfile.ZipFile(path) as package:
        return {
            name: hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        }


class DocumentInvariants(unittest.TestCase):
    def setUp(self):
        self.before = Document(SOURCE_PATH)
        self.after = Document(OUTPUT_PATH)

    def test_structure_is_preserved(self):
        self.assertEqual(len(self.before.paragraphs), len(self.after.paragraphs))
        self.assertEqual(len(self.before.inline_shapes), 28)
        self.assertEqual(len(self.after.inline_shapes), 28)
        self.assertEqual(len(self.before.tables), 12)
        self.assertEqual(len(self.after.tables), 12)
        self.assertEqual(
            [(len(table.rows), len(table.columns)) for table in self.before.tables],
            [(len(table.rows), len(table.columns)) for table in self.after.tables],
        )

    def test_headings_and_captions_are_preserved(self):
        protected_styles = {
            "Heading 2",
            "Heading 3",
            "Mushroomie Figure Caption",
            "Mushroomie Table Caption",
        }
        protected_text = lambda document: [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name in protected_styles
        ]
        self.assertEqual(protected_text(self.before), protected_text(self.after))

    def test_embedded_media_names_and_hashes_are_preserved(self):
        self.assertEqual(media_hashes(SOURCE_PATH), media_hashes(OUTPUT_PATH))

    def test_numeric_evidence_is_preserved(self):
        get_numbers = lambda document: re.findall(
            r"\b\d+(?:[.,]\d+)?(?:/\d+)?\b",
            document_text(document),
        )
        self.assertEqual(get_numbers(self.before), get_numbers(self.after))

    def test_campaign_plan_and_urls_are_preserved(self):
        get_stages = lambda document: re.findall(
            r"^➔ Giai đoạn .+", document_text(document), re.MULTILINE
        )
        get_activities = lambda document: re.findall(
            r"^Hoạt động \d+:.+", document_text(document), re.MULTILINE
        )
        get_urls = lambda document: re.findall(
            r"https?://[^\s,;)\]]+", document_text(document)
        )

        self.assertEqual(len(get_stages(self.before)), 4)
        self.assertEqual(len(get_activities(self.before)), 11)
        self.assertEqual(get_urls(self.before), ["https://mushroomie.io.vn/"])
        self.assertEqual(get_stages(self.before), get_stages(self.after))
        self.assertEqual(get_activities(self.before), get_activities(self.after))
        self.assertEqual(get_urls(self.before), get_urls(self.after))

    def test_required_rules(self):
        text = document_text(self.after)
        self.assertNotIn("landing page", text.lower())
        self.assertTrue(
            all(
                paragraph.paragraph_format.first_line_indent in (None, 0)
                for paragraph in all_paragraphs(self.after)
            )
        )
        keyword_table = self.after.tables[6]
        self.assertEqual(
            [keyword_table.cell(index, 0).text.strip() for index in range(1, 31)],
            [str(index) for index in range(1, 31)],
        )

    def test_rewrite_scope_and_application(self):
        self.assertEqual(len(PARAGRAPH_REWRITES), 66)
        self.assertEqual(len(TABLE_CELL_REWRITES), 19)
        before_text = document_text(self.before)
        after_text = document_text(self.after)

        for old_text, new_text in PARAGRAPH_REWRITES.items():
            self.assertIn(old_text, before_text)
            self.assertNotIn(old_text, after_text)
            self.assertIn(new_text, after_text)
        for old_text, new_text in TABLE_CELL_REWRITES.items():
            self.assertIn(old_text, before_text)
            self.assertNotIn(old_text, after_text)
            self.assertIn(new_text, after_text)

    def test_only_mapped_text_is_changed(self):
        changed_paragraphs = []
        matched_paragraphs = []
        for index, (before, after) in enumerate(
            zip(self.before.paragraphs, self.after.paragraphs)
        ):
            expected = PARAGRAPH_REWRITES.get(before.text, before.text)
            with self.subTest(paragraph=index):
                self.assertEqual(after.text, expected)
            if before.text in PARAGRAPH_REWRITES:
                matched_paragraphs.append(before.text)
            if before.text != after.text:
                changed_paragraphs.append(index)

        changed_cells = []
        matched_cells = []
        for table_index, (before_table, after_table) in enumerate(
            zip(self.before.tables, self.after.tables)
        ):
            for row_index, (before_row, after_row) in enumerate(
                zip(before_table.rows, after_table.rows)
            ):
                for cell_index, (before_cell, after_cell) in enumerate(
                    zip(before_row.cells, after_row.cells)
                ):
                    self.assertEqual(
                        len(before_cell.paragraphs), len(after_cell.paragraphs)
                    )
                    for paragraph_index, (before, after) in enumerate(
                        zip(before_cell.paragraphs, after_cell.paragraphs)
                    ):
                        expected = TABLE_CELL_REWRITES.get(before.text, before.text)
                        coordinate = (
                            table_index,
                            row_index,
                            cell_index,
                            paragraph_index,
                        )
                        with self.subTest(table_cell=coordinate):
                            self.assertEqual(after.text, expected)
                        if before.text in TABLE_CELL_REWRITES:
                            matched_cells.append(before.text)
                        if before.text != after.text:
                            changed_cells.append(coordinate)

        self.assertEqual(len(changed_paragraphs), len(PARAGRAPH_REWRITES))
        self.assertEqual(set(matched_paragraphs), set(PARAGRAPH_REWRITES))
        self.assertEqual(len(matched_paragraphs), len(PARAGRAPH_REWRITES))
        self.assertEqual(len(changed_cells), len(TABLE_CELL_REWRITES))
        self.assertEqual(set(matched_cells), set(TABLE_CELL_REWRITES))
        self.assertEqual(len(matched_cells), len(TABLE_CELL_REWRITES))

    def test_modality_markers_are_preserved_in_rewrites(self):
        modality_markers = ("cần", "nên", "có thể", "góp phần", "nếu")
        for old_text, new_text in (
            PARAGRAPH_REWRITES | TABLE_CELL_REWRITES
        ).items():
            for marker in modality_markers:
                if marker in old_text.lower():
                    with self.subTest(marker=marker, source=old_text[:80]):
                        self.assertIn(marker, new_text.lower())

        guarded_phrases = {
            "Từ góc độ nhận diện thương hiệu": "bố cục cần ưu tiên",
            "Trang chi tiết sản phẩm cần được đánh giá":
                "Trang chi tiết sản phẩm cần được đánh giá",
            "Đánh giá trải nghiệm di động":
                "thẻ sản phẩm nên hiển thị",
            "Bên cạnh điểm tổng quan": "việc đánh giá cần đi sâu",
            "Trong giai đoạn chạy quảng cáo":
                "quảng cáo tìm kiếm Google được xem là công cụ có thể",
            "Việc tối giản số trường bắt buộc": "góp phần hạn chế",
            "Để bảo đảm khả năng kiểm chứng":
                "Minh chứng nên được lưu theo tên từ khóa",
        }
        for source_prefix, required_phrase in guarded_phrases.items():
            matching_values = [
                new_text
                for old_text, new_text in PARAGRAPH_REWRITES.items()
                if old_text.startswith(source_prefix)
            ]
            with self.subTest(source_prefix=source_prefix):
                self.assertEqual(len(matching_values), 1)
                self.assertIn(required_phrase, matching_values[0])

    def test_high_risk_semantic_modality_is_preserved(self):
        paragraph_guards = {
            6: {
                "required": ("bố cục cần ưu tiên",),
                "forbidden": ("nhóm ưu tiên ảnh sản phẩm",),
            },
            29: {
                "required": ("Trang chi tiết sản phẩm cần được đánh giá",),
                "forbidden": ("Nhóm đánh giá trang chi tiết",),
            },
            62: {
                "required": (
                    "thẻ sản phẩm nên hiển thị",
                    "nút kêu gọi hành động cần có",
                    "nội dung không được tràn ngang",
                ),
                "forbidden": ("thẻ sản phẩm hiển thị theo",),
            },
            68: {
                "required": ("nhóm xây dựng quy chuẩn đăng sản phẩm",),
                "forbidden": ("nhóm cần xây dựng quy chuẩn",),
            },
            103: {
                "required": (
                    "việc đánh giá cần đi sâu",
                    "chỉ được ghi nhận khi có ảnh chụp",
                ),
                "forbidden": (
                    "nhóm phân tích từng nhóm tài nguyên",
                    "chỉ nên được ghi nhận",
                ),
            },
            109: {
                "required": ("Nhóm thực hiện báo cáo trực tiếp",),
                "forbidden": ("Nhóm cần thực hiện báo cáo",),
            },
            143: {
                "required": ("Minh chứng nên được lưu theo tên từ khóa",),
                "forbidden": ("Việc lưu minh chứng theo tên từ khóa giúp",),
            },
            148: {
                "required": ("Nhóm triển khai liên kết ngược",),
                "forbidden": ("Nhóm cần triển khai liên kết ngược",),
            },
            155: {
                "required": ("quảng cáo tìm kiếm Google được xem là công cụ có thể",),
                "forbidden": (
                    "nhóm sử dụng quảng cáo tìm kiếm Google",
                    "Các truy vấn có ý định mua hộp quà hoặc quà tặng handmade được dẫn",
                ),
            },
            164: {
                "required": ("nhóm đối chiếu", "Nhóm cũng rà soát Website"),
                "forbidden": ("nhóm cần đối chiếu",),
            },
        }
        for index, guard in paragraph_guards.items():
            text = self.after.paragraphs[index].text
            for required in guard["required"]:
                with self.subTest(paragraph=index, required=required):
                    self.assertIn(required, text)
            for forbidden in guard["forbidden"]:
                with self.subTest(paragraph=index, forbidden=forbidden):
                    self.assertNotIn(forbidden, text)

        negative_source = next(
            source
            for source in TABLE_CELL_REWRITES
            if source.startswith("Không bổ sung JavaScript")
        )
        negative_rewrite = TABLE_CELL_REWRITES[negative_source]
        self.assertIn("nếu không cần thiết", negative_rewrite)
        self.assertNotIn("nếu cần thiết", negative_rewrite)
