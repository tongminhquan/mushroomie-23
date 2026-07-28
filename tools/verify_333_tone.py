"""Final structural verification for the edited Mushroomie section 3.3.3.

The editor may change only approved prose.  This verifier therefore compares
the source and output DOCX at evidence and OOXML-layout level: each drawing's
relationship/size/crop, all table-grid and cell-width values, and protected
styles, settings, numbering, headers, footers, and relationship parts.
"""

from __future__ import annotations

import hashlib
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

try:
    from tools.sync_333_prose_to_asm import OUTPUT_PATH, SOURCE_PATH
except ModuleNotFoundError:  # Support direct execution from the tools directory.
    from sync_333_prose_to_asm import OUTPUT_PATH, SOURCE_PATH


NUMBER_PATTERN = re.compile(r"\b\d+(?:[.,]\d+)?(?:/\d+)?\b")
URL_PATTERN = re.compile(r"https?://[^\s,;)\]]+")
STAGE_PATTERN = re.compile(r"^➔ Giai đoạn .+", re.MULTILINE)
ACTIVITY_PATTERN = re.compile(r"^Hoạt động \d+:.+", re.MULTILINE)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


@dataclass(frozen=True)
class ParagraphRecord:
    location: str
    paragraph: Paragraph


@dataclass(frozen=True)
class DrawingSignature:
    location: str
    ordinal: int
    relationship_id: str
    target: str
    extent: tuple[str | None, str | None]
    crop: tuple[tuple[tuple[str, str], ...], ...]


@dataclass(frozen=True)
class TableSignature:
    width: tuple[str | None, str | None]
    indent: tuple[str | None, str | None]
    grid: tuple[tuple[str | None, str | None], ...]
    rows: tuple[tuple[tuple[str | None, str | None], ...], ...]


def iter_paragraphs(document: DocumentType) -> Iterator[ParagraphRecord]:
    """Yield body and table-cell paragraphs with stable human-readable paths."""
    for index, paragraph in enumerate(document.paragraphs):
        yield ParagraphRecord(f"P{index}", paragraph)
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield ParagraphRecord(
                        f"T{table_index}R{row_index}C{cell_index}P{paragraph_index}",
                        paragraph,
                    )


def document_text(document: DocumentType) -> str:
    return "\n".join(record.paragraph.text for record in iter_paragraphs(document))


def media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as package:
        return {
            name: hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        }


def canonical_xml(data: bytes) -> bytes:
    """Canonicalize XML so ZIP timestamps and XML attribute order do not matter."""
    return etree.tostring(etree.fromstring(data), method="c14n", with_comments=False)


def _strip_whitespace_nodes(element) -> None:
    """Remove serialisation-only indentation so semantic XML compares cleanly."""
    if element.text is not None and not element.text.strip():
        element.text = None
    for child in element:
        _strip_whitespace_nodes(child)
        if child.tail is not None and not child.tail.strip():
            child.tail = None


def protected_part_hashes(path: Path) -> dict[str, str]:
    """Hash canonical XML for structure/style parts that prose editing must not alter."""
    base_parts = {
        "word/styles.xml",
        "word/numbering.xml",
        "word/settings.xml",
        "word/_rels/document.xml.rels",
    }
    with zipfile.ZipFile(path) as package:
        part_names = set(package.namelist())
        protected = base_parts & part_names
        protected.update(
            name
            for name in part_names
            if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
            or re.fullmatch(r"word/_rels/(?:header|footer)\d+\.xml\.rels", name)
        )
        return {
            name: hashlib.sha256(canonical_xml(package.read(name))).hexdigest()
            for name in sorted(protected)
        }


def _attribute_signature(element, attribute: str = "w:w") -> tuple[str | None, str | None]:
    if element is None:
        return (None, None)
    return (element.get(qn("w:type")), element.get(qn(attribute)))


def drawing_signatures(document: DocumentType) -> list[DrawingSignature]:
    """Capture every drawing's document order, media relationship, geometry, and crop."""
    signatures: list[DrawingSignature] = []
    for record in iter_paragraphs(document):
        drawings = record.paragraph._p.xpath(".//w:drawing")
        for ordinal, drawing in enumerate(drawings):
            blips = drawing.xpath(".//a:blip")
            if len(blips) != 1:
                raise ValueError(f"Expected one a:blip at {record.location} drawing {ordinal}")
            relationship_id = blips[0].get(f"{{{R_NS}}}embed") or ""
            relationship = record.paragraph.part.rels.get(relationship_id)
            target = relationship.target_ref if relationship is not None else "<missing>"
            extents = drawing.xpath(".//wp:extent")
            extent = (
                extents[0].get("cx") if extents else None,
                extents[0].get("cy") if extents else None,
            )
            crop = tuple(
                tuple(sorted(source_rect.attrib.items()))
                for source_rect in drawing.xpath(".//a:srcRect")
            )
            signatures.append(
                DrawingSignature(record.location, ordinal, relationship_id, target, extent, crop)
            )
    return signatures


def table_signature(table) -> TableSignature:
    """Return all geometry values that determine an existing table's layout."""
    table_properties = table._tbl.tblPr
    width = _attribute_signature(table_properties.find(qn("w:tblW")))
    indent = _attribute_signature(table_properties.find(qn("w:tblInd")))
    grid = table._tbl.tblGrid
    grid_columns = tuple(
        (column.get(qn("w:type")), column.get(qn("w:w"))) for column in grid.gridCol_lst
    )
    rows: list[tuple[tuple[str | None, str | None], ...]] = []
    for row in table.rows:
        widths: list[tuple[str | None, str | None]] = []
        for cell in row.cells:
            cell_properties = cell._tc.tcPr
            widths.append(
                _attribute_signature(
                    cell_properties.find(qn("w:tcW")) if cell_properties is not None else None
                )
            )
        rows.append(tuple(widths))
    return TableSignature(width, indent, grid_columns, tuple(rows))


def _as_dxa(value: tuple[str | None, str | None]) -> int | None:
    type_value, width_value = value
    # Word omits w:type on gridCol, where dxa is implicit.
    if type_value not in (None, "dxa") or width_value is None:
        return None
    try:
        return int(width_value)
    except ValueError:
        return None


def validate_table_geometry(signature: TableSignature) -> list[str]:
    """Validate relationships between tblW, grid columns, and all tcW values."""
    errors: list[str] = []
    table_width = _as_dxa(signature.width)
    grid_widths = [_as_dxa(column) for column in signature.grid]
    if table_width is None:
        errors.append("tblW must be an explicit dxa value")
    if not signature.grid or any(width is None for width in grid_widths):
        errors.append("tblGrid must contain explicit dxa gridCol widths")
        return errors
    grid_total = sum(width for width in grid_widths if width is not None)
    if table_width is not None and grid_total != table_width:
        errors.append(f"tblGrid sum {grid_total} differs from tblW {table_width}")
    for row_index, row in enumerate(signature.rows):
        cell_widths = [_as_dxa(width) for width in row]
        if len(row) != len(signature.grid):
            errors.append(
                f"row {row_index} has {len(row)} cells; expected {len(signature.grid)} grid columns"
            )
        if any(width is None for width in cell_widths):
            errors.append(f"row {row_index} has non-dxa or missing tcW")
        elif sum(width for width in cell_widths if width is not None) != grid_total:
            errors.append(
                f"row {row_index} tcW sum {sum(cell_widths)} differs from grid sum {grid_total}"
            )
    return errors


def _xml_without_first_line_indent(paragraph: Paragraph) -> bytes:
    """Canonicalize paragraph properties while allowing the requested indent reset."""
    properties = paragraph._p.pPr
    if properties is None:
        return b""
    copied = etree.fromstring(properties.xml.encode("utf-8"))
    indent = copied.find(qn("w:ind"))
    if indent is not None:
        indent.attrib.pop(qn("w:firstLine"), None)
        indent.attrib.pop(qn("w:firstLineChars"), None)
        if not indent.attrib:
            copied.remove(indent)
    _strip_whitespace_nodes(copied)
    return etree.tostring(copied, method="c14n", with_comments=False)


def _run_properties(paragraph: Paragraph) -> list[bytes]:
    """Return direct run formatting only, not the allowed prose text values."""
    return [
        etree.tostring(etree.fromstring(run._r.rPr.xml.encode("utf-8")), method="c14n")
        if run._r.rPr is not None
        else b""
        for run in paragraph.runs
    ]


def _assert_equal(errors: list[str], label: str, before, after) -> None:
    if before != after:
        errors.append(f"{label}: source and output differ")


def verify(source: Path = SOURCE_PATH, output: Path = OUTPUT_PATH) -> list[str]:
    """Return final-QA failures; an empty list means the DOCX passed."""
    errors: list[str] = []
    if not source.is_file():
        return [f"Source DOCX not found: {source}"]
    if not output.is_file():
        return [f"Output DOCX not found: {output}"]

    before = Document(source)
    after = Document(output)
    before_text = document_text(before)
    after_text = document_text(after)

    _assert_equal(errors, "body paragraph count", len(before.paragraphs), len(after.paragraphs))
    _assert_equal(errors, "table count", len(before.tables), len(after.tables))
    _assert_equal(errors, "inline image count", len(before.inline_shapes), len(after.inline_shapes))
    if len(after.inline_shapes) != 28:
        errors.append(f"expected 28 inline images, found {len(after.inline_shapes)}")
    if len(after.tables) != 12:
        errors.append(f"expected 12 tables, found {len(after.tables)}")
    _assert_equal(
        errors,
        "table dimensions",
        [(len(table.rows), len(table.columns)) for table in before.tables],
        [(len(table.rows), len(table.columns)) for table in after.tables],
    )
    _assert_equal(errors, "embedded media hashes", media_hashes(source), media_hashes(output))
    _assert_equal(errors, "drawing signatures", drawing_signatures(before), drawing_signatures(after))
    _assert_equal(errors, "protected OOXML parts", protected_part_hashes(source), protected_part_hashes(output))

    _assert_equal(errors, "numeric evidence", NUMBER_PATTERN.findall(before_text), NUMBER_PATTERN.findall(after_text))
    _assert_equal(errors, "URLs", URL_PATTERN.findall(before_text), URL_PATTERN.findall(after_text))
    _assert_equal(errors, "stages", STAGE_PATTERN.findall(before_text), STAGE_PATTERN.findall(after_text))
    _assert_equal(errors, "activities", ACTIVITY_PATTERN.findall(before_text), ACTIVITY_PATTERN.findall(after_text))
    if len(STAGE_PATTERN.findall(after_text)) != 4:
        errors.append("expected 4 implementation stages")
    if len(ACTIVITY_PATTERN.findall(after_text)) != 11:
        errors.append("expected 11 implementation activities")
    if "landing page" in after_text.casefold():
        errors.append('forbidden phrase "landing page" is present')

    if len(after.tables) > 6:
        keyword_rows = [after.tables[6].cell(index, 0).text.strip() for index in range(1, 31)]
        if keyword_rows != [str(index) for index in range(1, 31)]:
            errors.append("Bảng 3.14 does not retain STT 1–30")
    else:
        errors.append("Bảng 3.14 cannot be checked because table index 6 is absent")

    protected_styles = {
        "Heading 2",
        "Heading 3",
        "Mushroomie Figure Caption",
        "Mushroomie Table Caption",
    }
    _assert_equal(
        errors,
        "headings and captions",
        [p.text for p in before.paragraphs if p.style.name in protected_styles],
        [p.text for p in after.paragraphs if p.style.name in protected_styles],
    )

    before_records = list(iter_paragraphs(before))
    after_records = list(iter_paragraphs(after))
    _assert_equal(errors, "total body/table paragraph count", len(before_records), len(after_records))
    for before_record, after_record in zip(before_records, after_records):
        if before_record.location != after_record.location:
            errors.append(f"paragraph path changed: {before_record.location}")
            continue
        if before_record.paragraph.style.name != after_record.paragraph.style.name:
            errors.append(f"paragraph style changed at {before_record.location}")
        if _xml_without_first_line_indent(before_record.paragraph) != _xml_without_first_line_indent(after_record.paragraph):
            errors.append(f"paragraph formatting changed at {before_record.location}")
        if _run_properties(before_record.paragraph) != _run_properties(after_record.paragraph):
            errors.append(f"run formatting changed at {before_record.location}")
        indent = after_record.paragraph.paragraph_format.first_line_indent
        if indent not in (None, 0):
            errors.append(f"first-line indent remains at {after_record.location}")

    for table_index, (before_table, after_table) in enumerate(zip(before.tables, after.tables)):
        before_signature = table_signature(before_table)
        after_signature = table_signature(after_table)
        _assert_equal(errors, f"table {table_index} geometry", before_signature, after_signature)
        for geometry_error in validate_table_geometry(after_signature):
            errors.append(f"table {table_index}: {geometry_error}")

    _assert_equal(
        errors,
        "section properties",
        [section._sectPr.xml for section in before.sections],
        [section._sectPr.xml for section in after.sections],
    )
    return errors


def main() -> int:
    errors = verify()
    if errors:
        print(f"FAIL: {len(errors)} final QA finding(s)")
        for error in errors:
            print(f"- {error}")
        return 1
    print("PASS: 4 stages; 11 activities; 28 inline images; 12 tables; Bảng 3.14 STT 1–30.")
    print("PASS: media/drawing targets, numeric evidence, URLs, headings/captions, protected OOXML, formatting, and exact table geometry are preserved.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
