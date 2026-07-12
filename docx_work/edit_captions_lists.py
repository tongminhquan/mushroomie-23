from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


TABLE_CAPTIONS = {
    5: ("Bảng 1.1", "Thông tin liên lạc của Mushroomie"),
    6: ("Bảng 1.2", "Danh mục sản phẩm chính của Mushroomie"),
    7: ("Bảng 1.3", "Danh mục sản phẩm phụ của Mushroomie"),
    8: ("Bảng 2.1", "Nền tảng hoạt động của Tiệm của Hamster"),
    9: ("Bảng 2.2", "Nền tảng hoạt động của Happy With Love Shop"),
    10: ("Bảng 2.3", "Phân tích điểm mạnh và điểm yếu của thương hiệu Mushroomie"),
    11: ("Bảng 2.4", "Phân tích cơ hội và thách thức của thương hiệu Mushroomie"),
    12: ("Bảng 2.5", "Phân tích mô hình STP"),
    13: ("Bảng 2.6", "Các cấp độ sản phẩm của Mushroomie"),
    14: ("Bảng 2.7", "Danh mục sản phẩm Mushroomie"),
    15: ("Bảng 2.8", "Bảng giá sản phẩm đề xuất của Mushroomie"),
    16: ("Bảng 2.9", "Chính sách giá linh hoạt của Mushroomie"),
    17: ("Bảng 2.10", "Kênh phân phối của Mushroomie"),
    18: ("Bảng 2.11", "Chiến lược xúc tiến của Mushroomie"),
    19: ("Bảng 2.12", "Tuyến nội dung truyền thông của Mushroomie"),
    20: ("Bảng 2.13", "Định hướng xúc tiến theo từng giai đoạn"),
}

CAPTION_RE = re.compile(r"^\s*Bảng\s*\d+(?:\.\d+)?\s*[:：.]", re.IGNORECASE)
FIGURE_RE = re.compile(r"^\s*Hình\s+(\d+(?:\.\d+)?)\s*[:：]\s*(.+?)\s*$", re.IGNORECASE)


def qname(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def paragraph_text(p: etree._Element) -> str:
    return " ".join("".join(p.xpath(".//w:t/text()", namespaces=NS)).split())


def clear_paragraph_text(p: etree._Element) -> None:
    for run in list(p.findall(qname("r"))):
        p.remove(run)
    for fld in list(p.findall(qname("fldSimple"))):
        p.remove(fld)
    for hyperlink in list(p.findall(qname("hyperlink"))):
        p.remove(hyperlink)


def set_paragraph_text(p: etree._Element, text: str) -> None:
    clear_paragraph_text(p)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)


def make_caption_paragraph(text: str) -> etree._Element:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "Caption")
    p_pr.append(style)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    p_pr.append(spacing)
    p.append(p_pr)
    set_paragraph_text(p, text)
    return p


def previous_text_paragraph(body_children: list[etree._Element], idx: int) -> etree._Element | None:
    for j in range(idx - 1, -1, -1):
        if etree.QName(body_children[j]).localname != "p":
            continue
        if paragraph_text(body_children[j]):
            return body_children[j]
    return None


def patch_table_captions(input_docx: Path, output_docx: Path) -> None:
    with zipfile.ZipFile(input_docx, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    body = root.find("w:body", NS)
    if body is None:
        raise RuntimeError("word/document.xml has no body")

    table_no = 0
    for child in list(body):
        if etree.QName(child).localname != "tbl":
            continue
        table_no += 1
        if table_no not in TABLE_CAPTIONS:
            continue

        label, title = TABLE_CAPTIONS[table_no]
        caption_text = f"{label}: {title}"
        children = list(body)
        table_idx = children.index(child)
        prev_p = previous_text_paragraph(children, table_idx)
        prev_text = paragraph_text(prev_p) if prev_p is not None else ""

        if prev_p is not None and re.match(r"^\s*Bảng\b", prev_text, re.IGNORECASE):
            # A broken caption like "Bảng :" or "Bảng 2:" is replaced; a full
            # existing "Bảng 2.3:" would be left alone by the caller's map.
            if not CAPTION_RE.match(prev_text) or label not in prev_text:
                set_paragraph_text(prev_p, caption_text)
        else:
            body.insert(table_idx, make_caption_paragraph(caption_text))

    with zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            if name == "word/document.xml":
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            zout.writestr(name, data)


def clear_table_keep_header(table) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def write_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text.isdigit() else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def canonical_caption(text: str) -> str:
    text = " ".join(text.split())
    text = re.sub(r"\s*[:：]\s*", ": ", text, count=1)
    return text


def fill_front_matter_tables(docx_path: Path, page_map: dict[str, int] | None = None) -> None:
    doc = Document(docx_path)
    page_map = page_map or {}
    canonical_page_map = {canonical_caption(k): v for k, v in page_map.items()}

    table_entries = [
        (label, title, str(canonical_page_map.get(canonical_caption(f"{label}: {title}"), "")))
        for label, title in TABLE_CAPTIONS.values()
    ]

    figure_entries = []
    seen = set()
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        m = FIGURE_RE.match(text)
        if not m:
            continue
        label = f"Hình {m.group(1)}"
        title = m.group(2).strip()
        full = f"{label}: {title}"
        if full in seen:
            continue
        seen.add(full)
        page = canonical_page_map.get(canonical_caption(full), "")
        figure_entries.append((label, title, str(page)))

    # Tables 3 and 4 are the existing "DANH MỤC BẢNG BIỂU" and
    # "DANH MỤC SƠ ĐỒ, HÌNH VẼ" front-matter tables in this document.
    for table_index, entries in [(2, table_entries), (3, figure_entries)]:
        table = doc.tables[table_index]
        clear_table_keep_header(table)
        header = table.rows[0].cells
        header[0].text = "Ký Hiệu"
        header[1].text = "Nội Dung"
        header[2].text = "Trang"
        for label, title, page in entries:
            row = table.add_row().cells
            write_cell(row[0], label)
            write_cell(row[1], title)
            write_cell(row[2], page)

    doc.save(docx_path)


def extract_pages_from_pdf(pdf_path: Path) -> dict[str, int]:
    import fitz

    page_map: dict[str, int] = {}
    doc = fitz.open(str(pdf_path))
    caption_patterns = []
    for label, title in TABLE_CAPTIONS.values():
        caption_patterns.append(f"{label}: {title}")

    for page_index in range(doc.page_count):
        text = " ".join(doc.load_page(page_index).get_text("text").split())
        for caption in caption_patterns:
            if caption in text and caption not in page_map:
                page_map[caption] = page_index + 1
        for match in re.finditer(r"Hình\s+(\d+(?:\.\d+)?)\s*[:：]\s*([^\\n\\r]+?)(?=(?:\\s{2,}| Hình\\s+\\d| Bảng\\s+\\d|$))", text):
            full = f"Hình {match.group(1)}: {match.group(2).strip()}"
            if full not in page_map:
                page_map[full] = page_index + 1
    return page_map


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--render-script", type=Path)
    parser.add_argument("--render-dir", type=Path)
    args = parser.parse_args()

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    temp_docx = args.output_docx.with_name(args.output_docx.stem + "_captions.docx")
    patch_table_captions(args.input_docx, temp_docx)
    shutil.copyfile(temp_docx, args.output_docx)
    fill_front_matter_tables(args.output_docx)

    if args.render_script and args.render_dir:
        args.render_dir.mkdir(parents=True, exist_ok=True)
        subprocess.run(
            [
                sys.executable,
                str(args.render_script),
                str(args.output_docx),
                "--output_dir",
                str(args.render_dir),
                "--emit_pdf",
            ],
            check=True,
        )
        pdf_path = args.render_dir / f"{args.output_docx.stem}.pdf"
        page_map = extract_pages_from_pdf(pdf_path)
        fill_front_matter_tables(args.output_docx, page_map)


if __name__ == "__main__":
    main()
